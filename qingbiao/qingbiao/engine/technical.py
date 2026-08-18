from __future__ import annotations

import json
import re
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from qingbiao.config import AREA_TOLERANCE_PCT, RESOURCES, TEXT_SIMILAR_THRESHOLD


def _load_json(name: str) -> dict[str, Any]:
    return json.loads((RESOURCES / name).read_text(encoding="utf-8"))


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("　".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"技术标暂支持 docx / pdf / txt，收到 {suffix}")


def split_paragraphs(text: str) -> list[str]:
    chunks = re.split(r"[\n\r]+|[。；;]", text)
    out = []
    for c in chunks:
        s = re.sub(r"\s+", "", c)
        if len(s) >= 18:
            out.append(s)
    return out


def _ratio(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a[:1200], b[:1200]).ratio()


def check_standards(text: str, structure_label: str) -> list[dict[str, Any]]:
    data = _load_json("standards.json")
    findings: list[dict[str, Any]] = []
    for item in data["superseded"]:
        if re.search(item["pattern"], text, flags=re.I):
            findings.append(
                {
                    "module": "技术标",
                    "category": "引用过期标准",
                    "severity": "高",
                    "detail": f"出现 {item['name']} 旧号，现行宜为 {item['replace']}",
                }
            )
    cited = set(re.findall(r"(?:GB/?T?|JGJ|GBJ)\s*[0-9]{4,5}\s*[-—－]\s*[0-9]{2,4}", text, flags=re.I))
    cited_norm = {re.sub(r"\s+", "", c).replace("—", "-").replace("－", "-").upper() for c in cited}
    needed = []
    for row in data["mandatory_current"]:
        applies = row["applies"]
        if "all" in applies or any(a in structure_label for a in applies):
            needed.append(row)
    for row in needed:
        code_key = re.sub(r"\s+", "", row["code"]).upper()
        if not any(code_key.split("-")[0] in c for c in cited_norm):
            findings.append(
                {
                    "module": "技术标",
                    "category": "未引用现行强制性通用规范",
                    "severity": "中",
                    "detail": f"概况结构类型为「{structure_label}」，建议核对是否引用 {row['code']}《{row['name']}》",
                }
            )
    return findings


def check_typos(text: str) -> list[dict[str, Any]]:
    data = _load_json("typos.json")
    findings: list[dict[str, Any]] = []
    for wrong, right in data["错别字"]:
        if wrong == right:
            continue
        if wrong in text:
            findings.append(
                {
                    "module": "技术标",
                    "category": "错别字",
                    "severity": "中",
                    "detail": f"出现「{wrong}」，疑应为「{right}」",
                }
            )
    for token in data["重复"]:
        if token in text:
            findings.append(
                {
                    "module": "技术标",
                    "category": "错别字",
                    "severity": "低",
                    "detail": f"出现重复用词「{token}」",
                }
            )
    return findings


def _match_structure(label: str) -> dict[str, Any] | None:
    data = _load_json("structure_types.json")
    for row in data.values():
        names = [row["label"], *row.get("aliases", [])]
        if any(n in label or label in n for n in names):
            return row
    return None


def check_structure_mismatch(text: str, structure_label: str) -> list[dict[str, Any]]:
    row = _match_structure(structure_label)
    if not row:
        return [
            {
                "module": "技术标",
                "category": "结构类型未识别",
                "severity": "低",
                "detail": f"概况填写为「{structure_label}」，未匹配内置结构类型词库，仅做弱校验",
            }
        ]
    findings: list[dict[str, Any]] = []
    expect_hits = sum(1 for w in row["expect"] if w in text)
    conflict_hits = [w for w in row["conflict"] if w in text]
    if expect_hits == 0:
        findings.append(
            {
                "module": "技术标",
                "category": "与结构类型不匹配",
                "severity": "高",
                "detail": f"概况为「{row['label']}」，正文几乎未出现 { '、'.join(row['expect'][:4]) } 等应有表述",
            }
        )
    if conflict_hits:
        findings.append(
            {
                "module": "技术标",
                "category": "与结构类型不匹配",
                "severity": "高",
                "detail": f"概况为「{row['label']}」，但正文出现冲突表述：{'、'.join(conflict_hits)}",
            }
        )
    return findings


def check_profile_consistency(text: str, profile: dict[str, Any]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    floors = str(profile.get("floors") or "").strip()
    if floors:
        nums = re.findall(r"(\d+)\s*层", text)
        if nums:
            distinct = {int(x) for x in nums if int(x) > 0}
            want = int(re.findall(r"\d+", floors)[0]) if re.findall(r"\d+", floors) else None
            if want and distinct and want not in distinct:
                findings.append(
                    {
                        "module": "技术标",
                        "category": "与基本概况不符",
                        "severity": "高",
                        "detail": f"概况层数为 {floors}，正文出现的层数有 {sorted(distinct)[:8]}",
                    }
                )
    area = profile.get("area")
    try:
        area_f = float(area) if area not in (None, "") else None
    except (TypeError, ValueError):
        area_f = None
    if area_f:
        areas = [float(x) for x in re.findall(r"(?:建筑面积|总建筑面积)\s*[:：]?\s*([\d.]+)", text)]
        for found in areas:
            if abs(found - area_f) / max(area_f, 1e-6) > AREA_TOLERANCE_PCT and abs(found - area_f) > 5:
                findings.append(
                    {
                        "module": "技术标",
                        "category": "与基本概况不符",
                        "severity": "中",
                        "detail": f"概况建筑面积 {area_f}，正文写到 {found}",
                    }
                )
                break
    return findings


def validate_one(text: str, profile: dict[str, Any]) -> list[dict[str, Any]]:
    label = str(profile.get("structure") or "")
    findings = []
    findings.extend(check_standards(text, label))
    findings.extend(check_typos(text))
    findings.extend(check_structure_mismatch(text, label))
    findings.extend(check_profile_consistency(text, profile))
    if len(text.strip()) < 80:
        findings.append(
            {
                "module": "技术标",
                "category": "文本提取不足",
                "severity": "高",
                "detail": "未能提取到足够正文（扫描件 PDF 可能无文字层）。请改用 Word 或可选中文字的 PDF。",
            }
        )
    return findings


def cross_similar(
    docs: list[dict[str, Any]],
    threshold: float = TEXT_SIMILAR_THRESHOLD,
) -> list[dict[str, Any]]:
    """docs: {bidder, paragraphs, text}"""
    findings: list[dict[str, Any]] = []
    for i in range(len(docs)):
        for j in range(i + 1, len(docs)):
            a, b = docs[i], docs[j]
            overall = _ratio(a.get("text") or "", b.get("text") or "")
            if overall >= threshold:
                findings.append(
                    {
                        "module": "技术标",
                        "category": "全文高度一致",
                        "severity": "高",
                        "bidder": f"{a['bidder']} / {b['bidder']}",
                        "detail": f"全文相似度为 {overall:.1%}，存在整体抄袭或同源模板风险",
                    }
                )
            pa, pb = a.get("paragraphs") or [], b.get("paragraphs") or []
            hits = 0
            samples: list[str] = []
            used = set()
            for x in pa:
                best = 0.0
                best_i = -1
                for yi, y in enumerate(pb):
                    if yi in used:
                        continue
                    r = _ratio(x, y)
                    if r > best:
                        best = r
                        best_i = yi
                if best >= threshold and best_i >= 0:
                    hits += 1
                    used.add(best_i)
                    if len(samples) < 3:
                        samples.append(x[:80] + ("…" if len(x) > 80 else ""))
            if hits >= 3 or (hits and hits / max(len(pa), 1) >= 0.15):
                findings.append(
                    {
                        "module": "技术标",
                        "category": "段落内容高度一致",
                        "severity": "高",
                        "bidder": f"{a['bidder']} / {b['bidder']}",
                        "detail": f"高度相似段落约 {hits} 处。示例：「{'；'.join(samples)}」",
                    }
                )
    return findings
