from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from qingbiao.config import DATA_DIR, ensure_dirs


def _set_run_font(run, size: int = 12, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color is not None:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def _add_para(doc: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, size=10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val or ""))
            _set_run_font(run, size=10)


def _sev_color(sev: str) -> RGBColor:
    if sev == "高":
        return RGBColor(0xB0, 0x18, 0x18)
    if sev == "中":
        return RGBColor(0xB3, 0x6B, 0x00)
    return RGBColor(0x33, 0x33, 0x33)


def build_report(project: dict[str, Any], results: dict[str, Any], dest: Path | None = None) -> Path:
    ensure_dirs()
    dest = dest or (DATA_DIR / "清标报告.docx")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Pt(595)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("建设工程投标文件清标报告")
    _set_run_font(run, size=22, bold=True)
    _add_para(doc, f"项目名称：{project.get('name') or '（未填写）'}")
    _add_para(doc, f"清标时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _add_para(
        doc,
        "结构类型：{structure}　层数：{floors}　建筑面积：{area}".format(
            structure=project.get("structure") or "-",
            floors=project.get("floors") or "-",
            area=project.get("area") or "-",
        ),
    )
    _add_para(doc, "本报告由本地清标助手离线生成，仅供清标复核使用，不替代评标委员会结论。", size=10)

    eco = results.get("economic") or {}
    tech = results.get("technical") or {}
    meta = results.get("metadata") or []

    _add_heading(doc, "一、清标范围", 1)
    _add_para(doc, f"经济标投标人：{'、'.join(eco.get('bidders') or []) or '无'}")
    _add_para(doc, f"最高投标限价文件：{eco.get('limit_file') or '未上传'}")
    _add_para(doc, f"技术标投标人：{'、'.join(tech.get('bidders') or []) or '无'}")

    _add_heading(doc, "二、经济标问题", 1)
    eco_findings = eco.get("findings") or []
    if not eco_findings:
        _add_para(doc, "经济标未发现需写入报告的问题（或尚未分析）。")
    else:
        rows = []
        for f in eco_findings:
            rows.append(
                [
                    f.get("bidder") or "",
                    f.get("category") or "",
                    f.get("severity") or "",
                    f.get("item_code") or "",
                    f.get("item_name") or "",
                    f.get("detail") or "",
                ]
            )
        _table(doc, ["谁家", "什么问题", "程度", "项目编码", "项目名称", "说明"], rows)

    _add_heading(doc, "三、技术标单份验证", 1)
    single = tech.get("single") or []
    if not single:
        _add_para(doc, "技术标尚未做单份验证，或未发现问题。")
    else:
        rows = []
        for f in single:
            rows.append(
                [
                    f.get("bidder") or "",
                    f.get("category") or "",
                    f.get("severity") or "",
                    f.get("detail") or "",
                ]
            )
        _table(doc, ["谁家", "什么问题", "程度", "说明"], rows)

    _add_heading(doc, "四、技术标横向对比", 1)
    cross = tech.get("cross") or []
    if not cross:
        _add_para(doc, "技术标横向对比未发现高度一致片段（或尚未分析）。")
    else:
        rows = [[f.get("bidder") or "", f.get("category") or "", f.get("severity") or "", f.get("detail") or ""] for f in cross]
        _table(doc, ["谁家", "什么问题", "程度", "说明"], rows)

    _add_heading(doc, "五、文件属性（同一账号 / 同一计算机）", 1)
    if not meta:
        _add_para(doc, "文件属性比对未发现相同作者、修改者或创建环境。")
    else:
        rows = [
            [
                " / ".join(f.get("bidders") or []),
                f.get("category") or "",
                f.get("severity") or "",
                f.get("detail") or "",
            ]
            for f in meta
        ]
        _table(doc, ["谁家", "什么问题", "程度", "相同属性"], rows)

    _add_heading(doc, "六、汇总", 1)
    all_f = list(eco_findings) + list(single) + list(cross) + list(meta)
    high = sum(1 for f in all_f if f.get("severity") == "高")
    mid = sum(1 for f in all_f if f.get("severity") == "中")
    _add_para(doc, f"合计写入问题 {len(all_f)} 条，其中高 {high} 条、中 {mid} 条。")
    _add_para(
        doc,
        "建议：对「多家单价相同」「全文/段落高度一致」「疑似同一账号」优先提交评标委员会核实；"
        "对「超过最高投标限价」「引用过期标准」「与结构类型不匹配」按招标文件废标/澄清条款处理。",
    )

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest
