from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook

from suanli_budget.config import DATA_DIR, ensure_dirs


def _font(run, size: int = 12, bold: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def yuan(v: float) -> str:
    return f"{v:,.2f}"


def export_excel(budget: dict[str, Any], dest: Path | None = None) -> Path:
    ensure_dirs()
    dest = dest or (DATA_DIR / "算力中心预算表.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "投资估算"
    meta = budget["meta"]
    scale = budget["scale"]
    tot = budget["totals"]
    ws.append(["项目", meta["project"], "地点", meta.get("location") or ""])
    ws.append(["GPU", meta["gpu_name"], "制冷", "液冷" if meta["cooling"] == "liquid" else "风冷"])
    ws.append(["采购卡数", scale["gpu_buy"], "节点", scale["nodes"], "FP16 PFLOPS", scale["pflops_fp16"]])
    ws.append(["IT功率kW", scale["it_kw"], "进线功率kW", scale["facility_kw"], "PUE", scale["pue"]])
    ws.append([])
    ws.append(["分项", "规格", "数量", "单位", "单价(元)", "合价(元)"])
    for sec in budget["sections"]:
        ws.append([sec["title"], "", "", "", "", sec["subtotal"]])
        for r in sec["rows"]:
            ws.append([r["name"], r.get("spec") or "", r["qty"], r["unit"], r["price"], r["amount"]])
    ws.append([])
    ws.append(["预备费", f"{tot['contingency_pct']}%", "", "", "", tot["contingency"]])
    ws.append(["不含税合计", "", "", "", "", tot["pretax"]])
    ws.append(["增值税", f"{tot['vat_pct']}%", "", "", "", tot["vat"]])
    ws.append(["含税总投资", "", "", "", "", tot["total"]])
    ws.append(["单卡含税造价", "", "", "", "", tot["per_gpu"]])
    ws.append(["每PFLOPS含税造价", "", "", "", "", tot["per_pflops"]])
    ws.append(["年电费（运维，不含在总投资）", "", "", "", "", tot["opex_power_year"]])
    ws.append([])
    ws.append(["说明", meta.get("price_note") or ""])
    dest.parent.mkdir(parents=True, exist_ok=True)
    wb.save(dest)
    return dest


def export_docx(budget: dict[str, Any], dest: Path | None = None) -> Path:
    ensure_dirs()
    dest = dest or (DATA_DIR / "算力中心预算书.docx")
    meta, scale, tot = budget["meta"], budget["scale"], budget["totals"]
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("算力中心投资估算预算书")
    _font(r, 22, True)
    p = doc.add_paragraph()
    r = p.add_run(f"项目：{meta['project']}　　地点：{meta.get('location') or '（未填）'}　　编制日期：{datetime.now():%Y-%m-%d}")
    _font(r, 11)
    doc.add_heading("一、编制说明", 1)
    p = doc.add_paragraph()
    r = p.add_run(
        "本预算由本地「算力中心预算编制」根据输入的算力目标、GPU/CPU 等基础数据和单价汇总。"
        "默认单价为 2026 年公开渠道参考价，成交价随集采批次变化，请以询价覆盖。"
        "整机价模式下 CPU/内存默认已含在服务器内。增值税按设备及建安合计计列，实际抵扣以财务为准。"
        "年电费单列，不计入建设总投资。"
    )
    _font(r, 11)
    doc.add_heading("二、建设规模", 1)
    for line in [
        f"GPU：{meta['gpu_name']}",
        f"目标方式：{'按卡数' if scale['mode']=='count' else '按算力(FP16 PFLOPS)'}；需求卡 {scale['want_gpu']}，按整机采购 {scale['gpu_buy']} 卡 / {scale['nodes']} 台。",
        f"标称 FP16 算力约 {scale['pflops_fp16']} PFLOPS（按单卡 {scale['tflops_per_card']} TFLOPS 线性折算，未计稀疏/集群效率）。",
        f"IT 功率约 {scale['it_kw']} kW，按 PUE {scale['pue']} 进线约 {scale['facility_kw']} kW；机柜 {scale['racks']} 套，面积约 {scale['area_m2']} ㎡；制冷：{'液冷' if meta['cooling']=='liquid' else '风冷'}。",
    ]:
        p = doc.add_paragraph(line)
        for run in p.runs:
            _font(run, 11)
    doc.add_heading("三、投资估算表", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["分项/名称", "规格", "数量", "单位", "单价(元)", "合价(元)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _font(run, 10, True)
    def add_row(vals: list[Any], bold: bool = False) -> None:
        cells = table.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            _font(run, 10, bold)

    for sec in budget["sections"]:
        add_row([sec["title"], "", "", "", "", yuan(sec["subtotal"])], True)
        for r in sec["rows"]:
            add_row([r["name"], r.get("spec") or "", r["qty"], r["unit"], yuan(r["price"]), yuan(r["amount"])])
    add_row(["预备费", f"{tot['contingency_pct']}%", "", "", "", yuan(tot["contingency"])], True)
    add_row(["不含税合计", "", "", "", "", yuan(tot["pretax"])], True)
    add_row(["增值税", f"{tot['vat_pct']}%", "", "", "", yuan(tot["vat"])], True)
    add_row(["含税总投资", "", "", "", "", yuan(tot["total"])], True)

    doc.add_heading("四、技术经济指标", 1)
    for line in [
        f"含税总投资 {yuan(tot['total'])} 元（约 {tot['total']/10000:,.2f} 万元）。",
        f"单卡含税造价 {yuan(tot['per_gpu'])} 元；每 PFLOPS（FP16 标称）含税造价 {yuan(tot['per_pflops'])} 元。",
        f"年电费约 {yuan(tot['opex_power_year'])} 元（按进线功率×年利用小时×电价，不含人工）。",
    ]:
        p = doc.add_paragraph(line)
        for run in p.runs:
            _font(run, 11)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest
