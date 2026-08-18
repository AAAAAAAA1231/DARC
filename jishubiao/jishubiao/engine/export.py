from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from jishubiao.config import DATA_DIR, ensure_dirs


def _run(paragraph, text: str, size: int = 12, bold: bool = False, center: bool = False) -> None:
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def export_docx(doc: dict[str, Any], dest: Path | None = None) -> Path:
    ensure_dirs()
    name = doc["project"]["name"]
    dest = dest or (DATA_DIR / f"{name}-技术标.docx")
    d = Document()
    for sec in d.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.6)
    p = d.add_paragraph()
    _run(p, doc["project"]["name"], 22, True, True)
    p = d.add_paragraph()
    _run(p, "施 工 组 织 设 计（技 术 标）", 18, True, True)
    p = d.add_paragraph()
    _run(
        p,
        f"投标人：{doc['project']['bidder']}    编制日期：{datetime.now():%Y年%m月%d日}",
        12,
        False,
        True,
    )
    p = d.add_paragraph()
    _run(p, "（本机生成初稿，须技术负责人审核后使用）", 10, False, True)

    d.add_heading("目录", level=1)
    for i, title in enumerate(doc["toc"], 1):
        p = d.add_paragraph()
        _run(p, f"{i}. {title}", 12)

    d.add_heading("引用规范一览", level=1)
    table = d.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "规范号", "名称", "类别"
    for row in doc["codes"]:
        cells = table.add_row().cells
        cells[0].text = row["code"]
        cells[1].text = row["name"]
        cells[2].text = row.get("kind") or ""

    for ch in doc["chapters"]:
        d.add_heading(ch["title"], level=1)
        for sec in ch["sections"]:
            d.add_heading(sec["heading"], level=2)
            for para in str(sec["body"]).split("\n"):
                if not para.strip():
                    continue
                p = d.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Cm(0.74)
                _run(p, para.strip(), 12)
            if sec.get("codes"):
                p = d.add_paragraph()
                _run(p, "主要依据：" + "、".join(sec["codes"]), 10)

    d.add_heading("使用声明", level=1)
    for w in doc.get("warnings") or []:
        p = d.add_paragraph()
        _run(p, w, 10)

    dest.parent.mkdir(parents=True, exist_ok=True)
    d.save(dest)
    return dest


def export_markdown(doc: dict[str, Any]) -> str:
    lines = [
        f"# {doc['project']['name']} 施工组织设计（技术标）",
        "",
        f"- 地点：{doc['project']['location']}",
        f"- 专业：{doc['project']['specialty']}　结构：{doc['project']['structure']}",
        f"- 投标人：{doc['project']['bidder']}",
        "",
        "## 引用规范一览",
        "",
        "| 规范号 | 名称 | 类别 |",
        "| --- | --- | --- |",
    ]
    for row in doc["codes"]:
        lines.append(f"| {row['code']} | {row['name']} | {row.get('kind') or ''} |")
    lines.append("")
    for ch in doc["chapters"]:
        lines.append(f"## {ch['title']}")
        lines.append("")
        for sec in ch["sections"]:
            lines.append(f"### {sec['heading']}")
            lines.append("")
            lines.append(sec["body"])
            lines.append("")
            if sec.get("codes"):
                lines.append("主要依据：" + "、".join(sec["codes"]))
                lines.append("")
    lines.append("## 使用声明")
    lines.append("")
    for w in doc.get("warnings") or []:
        lines.append(f"- {w}")
    return "\n".join(lines)
